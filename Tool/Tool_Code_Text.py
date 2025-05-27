import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

selected_files = []

def choose_files():
    files = filedialog.askopenfilenames(filetypes=[("All files", "*.*")])
    if files:
        selected_files.extend(files)
        file_paths.set("\n".join(selected_files))  # Cập nhật danh sách hiển thị

def clear_selection():
    global selected_files
    selected_files = []
    file_paths.set("")

def convert_files():
    if not selected_files:
        result_label.config(text="Vui lòng chọn các tệp trước!")
        return
    
    output_file = "output.docx"
    doc = Document()

    for file_path in selected_files:
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
            doc.add_paragraph(f"--- {file_path.split('/')[-1]} ---")
            doc.add_paragraph(content)
            doc.add_paragraph("\n")
        except Exception as e:
            doc.add_paragraph(f"Lỗi khi đọc {file_path}: {e}")

    doc.save(output_file)
    result_label.config(text=f"Tạo thành công: {output_file}")

def show_help():
    help_text = """Hướng dẫn sử dụng:
1. Nhấn 'Chọn tệp' để chọn nhiều file code Python.
2. Nếu muốn chọn lại, nhấn 'Xóa danh sách'.
3. Khi đã chọn đủ file, nhấn 'Chuyển đổi' để tạo file .docx.
4. File output.docx sẽ chứa nội dung của tất cả các file đã chọn.
"""
    messagebox.showinfo("Hướng dẫn sử dụng", help_text)

# Tạo UI với Tkinter
root = tk.Tk()
root.title("Chuyển File Code Sang .docx")

file_paths = tk.StringVar()

tk.Label(root, text="Chọn các tệp code:").pack()
tk.Button(root, text="Chọn tệp", command=choose_files).pack()
tk.Button(root, text="Xóa danh sách", command=clear_selection).pack()
tk.Label(root, textvariable=file_paths, wraplength=400, justify="left").pack()
tk.Button(root, text="Chuyển đổi", command=convert_files).pack()
result_label = tk.Label(root, text="")
result_label.pack()

tk.Button(root, text="Hướng dẫn sử dụng", command=show_help).pack()  # Thêm nút hướng dẫn

root.mainloop()
